import datetime
import os

import jieba
import moviepy.editor as mp
import numpy as np
import pandas as pd
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm
from google.cloud import speech, speech_v1p1beta1, storage
from nltk.tokenize import word_tokenize
from wordcloud import WordCloud

jieba.set_dictionary('dict.txt.big.txt')


def upload_blob(bucket_name, source_file_name, destination_blob_name, json_os):
    """
    上傳檔案到bucket
    bucket_name = "我們的bucket名稱"
    source_file_name = "我們要上傳檔案的所在「路徑」"
    destination_blob_name = "上傳storage後的名稱要叫做什麼"
    """

    # 環境設定
    os.environ['GOOGLE_APPLICATION_CREDENTIALS'] = json_os

    from google.cloud import storage

    storage_client = storage.Client.from_service_account_json(json_os)
    bucket = storage_client.bucket(bucket_name)
    blob = bucket.blob(destination_blob_name)

    blob.upload_from_filename(source_file_name)

    print("檔案 {} 已經被上傳到 {}.".format(source_file_name, bucket_name))


def translate_chinese(df, lan="zh_TW"):
    '''
    逐字辨識
    '''
    from google.cloud import translate_v2 as translate

    client = translate.Client()
    for i in range(0, len(df['斷字'])):
        response = client.translate(df['斷字'][i], target_language=lan)
        df['斷字'][i] = response["translatedText"]
    return df


def punctuation_auto_generate_subtitle_preprocessing(response, word_table):
    # transcript_list = []
    # for i in response.results:
    #     transcript = i.alternatives[0].transcript
    #     transcript_list.append(transcript)

    transcript_list = []
    # 結果呈現
    for result in response:
        transcript_list.append(result.alternatives[0].transcript)  # 辨識結果
        # print(result.alternatives[0].confidence) # 信心水準

    # 將transcript_list變成字串，爲後續的無句讀轉句讀做準備
    transcript_list_word = ''.join(transcript_list)

    # 因爲英文單子在word_table裡面算一個單字，所以要將table裡面的英文單字分開

    import re

    can = pd.DataFrame()

    for alpha in range(0, len(word_table)):
        autotemp = word_table[alpha : alpha + 1]
        line = re.sub("[^A-Za-z\d+]", "", autotemp['word'].values[0].strip())

        # 偵測有無英文
        if line != '':
            print(line)
            word_table_tmp = word_table[word_table['word'] == line]
            word_table_tmp = word_table_tmp[word_table_tmp.index == alpha]

            # 在到英文的row field
            for eng_word in word_table_tmp['word']:
                # 將其for迴圈拆解
                for word in eng_word:
                    print(word)
                    word_table_tmp['word'] = word
                    can = pd.concat([can, word_table_tmp], axis=0)

    can_index = list(set(can.index.to_list()))
    # word_table = word_table.copy()
    for i in range(0, len(can_index)):
        less_df = word_table[word_table.index < can_index[i]]
        can_tmp = can[can.index == can_index[i]]
        greater_df = word_table[word_table.index > can_index[i]]
        word_table = pd.concat([less_df, can_tmp, greater_df], axis=0)

    # 將裡面的英文轉換成小寫，以便對照word_table裡面的字詞
    transcript_list_word = transcript_list_word.lower()

    # 找出句讀的index
    index = -1
    index_list = []
    word_list = []
    for word in transcript_list_word:
        index += 1
        if re.search('，|？|。', word):
            index_list.append(index)
            word_list.append(word)

    pun_df = pd.DataFrame({'index': index_list, 'pun': word_list})

    pun_df['diff'] = range(0, len(pun_df))
    pun_df['index'] = pun_df['index'] - pun_df['diff']
    word_table = word_table.reset_index().drop(columns='index')

    # 使用for迴圈完成新的word_table，以便於我們後續切分字幕
    new_word_table = pd.DataFrame()

    for i in range(0, len(pun_df)):
        if i == 0:
            autotemp = pun_df[i : i + 1]
            temp = word_table[0 : autotemp['index'].values[0]]
            punctuation_df = temp.iloc[len(temp) - 1 : len(temp)].copy()
            pun_tmp = autotemp['pun'].values[0]
            punctuation_df['word'] = pun_tmp
            full_df = pd.concat([temp, punctuation_df], axis=0)

        else:
            autotemp = pun_df[i - 1 : i + 1]
            start = autotemp['index'].iloc[0]  # - minus +1 -2 # 4--> -2
            end = autotemp['index'].iloc[1]
            temp = word_table[start:end]
            punctuation_df = temp.iloc[len(temp) - 1 : len(temp)].copy()
            pun_tmp = autotemp.iloc[len(autotemp) - 1]['pun']
            punctuation_df['word'] = pun_tmp
            full_df = pd.concat([temp, punctuation_df], axis=0)

        new_word_table = pd.concat([new_word_table, full_df], axis=0)

    new_word_table['cut_para'] = np.where(
        new_word_table['word'] == '，',
        'cut',
        np.where(new_word_table['word'] == '？', 'cut', 'no'),
    )

    new_word_table = new_word_table.reset_index().drop(columns='index')

    # # 將有「cut」的index抓出來，方便我們切段落
    auto_select_range = new_word_table[
        new_word_table['cut_para'] == 'cut'
    ].index

    # 因爲是句讀後才要放
    auto_select_range = auto_select_range + 1

    return new_word_table, auto_select_range


def punctuation_auto_generate_subtitle_second_adj(new_word_table, sec):
    cut_point = new_word_table[new_word_table['cut_para'] == 'cut'].index
    #  = 5
    for i in range(0, len(cut_point)):
        if i == 0:
            new_word_table_tmp = new_word_table[0 : cut_point[i]]
            end = new_word_table_tmp.iloc[len(new_word_table_tmp) - 1][
                'end_time'
            ]
            start = new_word_table_tmp.iloc[0]['start_time']

            # 如果句讀之間的區間大於sec秒，就要cut
            if end - start >= sec:
                new_word_table_tmp['cut_para'] = np.where(
                    (
                        new_word_table_tmp['start_time_floor']
                        != new_word_table_tmp['end_time_floor']
                    )
                    & (new_word_table_tmp['end_time_floor'] % sec == 0),
                    'cut',
                    new_word_table_tmp['cut_para'],
                )

        elif i != len(cut_point) - 1:
            new_word_table_tmp = new_word_table[
                cut_point[i - 1] : cut_point[i]
            ]
            end = new_word_table_tmp.iloc[len(new_word_table_tmp) - 1][
                'end_time'
            ]
            start = new_word_table_tmp.iloc[0]['start_time']

            # new_word_table_tmp也會影響到new_word_table
            if end - start >= sec:
                # print(auto_select_list)
                new_word_table_tmp['cut_para'] = np.where(
                    (
                        new_word_table_tmp['start_time_floor']
                        != new_word_table_tmp['end_time_floor']
                    )
                    & (new_word_table_tmp['end_time_floor'] % sec == 0),
                    'cut',
                    new_word_table_tmp['cut_para'],
                )

        # 最後一個
        else:
            new_word_table_tmp = new_word_table[
                cut_point[i - 1] : len(new_word_table)
            ]
            end = new_word_table_tmp.iloc[len(new_word_table_tmp) - 1][
                'end_time'
            ]
            start = new_word_table_tmp.iloc[0]['start_time']

            # new_word_table_tmp也會影響到new_word_table
            if end - start >= sec:
                print('end')
                # print(auto_select_list)
                new_word_table_tmp['cut_para'] = np.where(
                    (
                        new_word_table_tmp['start_time_floor']
                        != new_word_table_tmp['end_time_floor']
                    )
                    & (new_word_table_tmp['end_time_floor'] % sec == 0),
                    'cut',
                    new_word_table_tmp['cut_para'],
                )

    # 重新切分index
    auto_select_range = new_word_table[
        new_word_table['cut_para'] == 'cut'
    ].index
    auto_select_range = auto_select_range + 1
    return new_word_table, auto_select_range


def auto_generate_subtitle(word_table, auto_select_range, video):
    auto_select_list = []
    for i in range(0, len(auto_select_range)):
        # 先抓第一個斷句前與後的文字 [0 ~ 11 ~ 17]
        if i == 0:
            # 開始時間
            autotemp = auto_select_range[i]
            temp = word_table[
                (word_table.index >= 0) & (word_table.index < autotemp)
            ]

            # 將該範圍的字詞組合成句子
            word_cut = ''.join(temp.word.tolist())

            # 演練題 - 我如何將 word_cut 這句的起始時間與結束時間抓出來?
            sentence_start_time = temp['start_time'].loc[0]
            sentence_end_time = temp['end_time'].loc[len(temp) - 1]

            '''
                把時間轉換成錄音檔的格式
                因為如果秒數剛好是整數的話後面就不會有微秒
                但錄音檔就算沒有微秒也要顯示000
                因此增加了下述判斷式
            '''

            # 影片特製秒數轉換
            sentence_start_time, sentence_end_time = srt_sec_format(
                sentence_start_time, sentence_end_time
            )

            auto_select_list.append(
                [word_cut, sentence_start_time, sentence_end_time]
            )

        # 第二個後的斷句，僅抓後面的斷字區間 [17 ~25], [25 ~ 32]...
        else:  # i > 0
            # 如果不是最後一個數字，就執行
            # 因為 [208]後面沒有數字了
            if i != len(auto_select_range) - 1:
                # 一次抓兩個出來，例如: [17 ~ 25]
                autotemp = auto_select_range[i - 1 : i + 1]
                temp = word_table[
                    (word_table.index >= autotemp[0])
                    & (word_table.index < autotemp[1])
                ]

                word_cut = ''.join(temp.word.tolist())
                sentence_start_time = temp['start_time'].loc[autotemp[0]]
                sentence_end_time = temp['end_time'].loc[autotemp[1] - 1]

                # 影片特製秒數轉換
                sentence_start_time, sentence_end_time = srt_sec_format(
                    sentence_start_time, sentence_end_time
                )

                auto_select_list.append(
                    [word_cut, sentence_start_time, sentence_end_time]
                )

            # 如果是最後一個數字，就執行else
            else:
                print('last one')
                autotemp = auto_select_range[i]
                temp = word_table[word_table.index >= autotemp]

                word_cut = ''.join(temp.word.tolist())
                sentence_start_time = temp['start_time'].loc[autotemp]
                sentence_end_time = temp['end_time'].loc[
                    autotemp + len(temp) - 1
                ]

                # 影片特製秒數轉換
                sentence_start_time, sentence_end_time = srt_sec_format(
                    sentence_start_time, sentence_end_time
                )

                auto_select_list.append(
                    [word_cut, sentence_start_time, sentence_end_time]
                )
    # srt_list完成字幕list
    srt_list = []
    i = 1

    # 透過迴圈，逐行新增srt形式字幕
    for write in auto_select_list:
        srt_list.append(
            str(i)
            + '\n'
            + write[1]
            + ' --> '
            + write[2]
            + '\n'
            + write[0]
            + '\n\n'
        )
        i += 1

    with open(video + '.srt', 'w', encoding='UTF-8') as f:
        f.write(''.join(srt_list))

    return srt_list


def auto_generate_subtitle_preprocessing(response, sec_split):
    # transcript_list = []
    # for i in response.results:
    #     transcript = i.alternatives[0].transcript
    #     transcript_list.append(transcript)

    transcript_list = []
    # 結果呈現
    for result in response:
        transcript_list.append(result.alternatives[0].transcript)  # 辨識結果
        # print(result.alternatives[0].confidence) # 信心水準

    ## 透過for迴圈，逐一將每個字的中文解字、始末時間放入清單`word_list`中
    word_list = []
    # 使用迴圈，取出每一段辨識內容
    for i in response:  # .results:
        # 使用迴圈，將選取逐字選取
        for word in i.alternatives[0].words:
            # 抓取文字
            word_ = word.word
            # 抓取開始時間
            s_time = round(
                word.start_time.seconds + word.start_time.microseconds * 1e-6,
                3,
            )
            # 抓取結束時間
            e_time = round(
                word.end_time.seconds + word.end_time.microseconds * 1e-6, 3
            )
            # 將結果一並添加至word_list中
            word_list.append([word_, s_time, e_time])

    word_table = pd.DataFrame(
        word_list, columns=['word', 'start_time', 'end_time']
    )

    # 開始時間整數
    word_table['start_time_floor'] = np.floor(word_table['start_time'])
    word_table

    # 結束時間整數
    word_table['end_time_floor'] = np.floor(word_table['end_time'])
    word_table

    word_table['cut_para'] = np.where(
        (word_table['start_time_floor'] != word_table['end_time_floor'])
        & (word_table['end_time_floor'] % sec_split == 0),
        'cut',
        'no',
    )

    word_table.loc[word_table.index < 5, 'cut_para'] = 'no'

    # 將有「cut」的index抓出來，方便我們切段落
    auto_select_range = word_table[word_table['cut_para'] == 'cut'].index

    return word_table, auto_select_range


def srt_sec_format(sentence_start_time, sentence_end_time):
    # 開始時間處理
    if np.floor(sentence_start_time) != sentence_start_time:
        # 情況一:該秒數含有小數點調整方法
        sentence_start_time = (
            '0' + str(datetime.timedelta(seconds=sentence_start_time))[:-3]
        )
    else:
        # 情況二:該秒數不含小數點調整方法
        sentence_start_time = (
            '0' + str(datetime.timedelta(seconds=sentence_start_time)) + '.000'
        )

    # 結束時間處理
    if np.floor(sentence_end_time) != sentence_end_time:
        # 情況一:該秒數含有小數點調整方法
        sentence_end_time = (
            '0' + str(datetime.timedelta(seconds=sentence_end_time))[:-3]
        )
    else:
        # 情況二:該秒數不含小數點調整方法
        sentence_end_time = (
            '0' + str(datetime.timedelta(seconds=sentence_end_time)) + '.000'
        )

    return sentence_start_time, sentence_end_time


def video2text_basic(video_path):
    clip = mp.VideoFileClip(video_path)

    clip.audio.write_audiofile(
        video_path.split('.')[0] + '.wav',
        fps=48000,  # 頻率
        ffmpeg_params=["-ac", "1"],
    )  # 編碼方式

    print(video_path + '.wav' + ' done!')


def move_file(dectect_name, folder_name):
    '''
    dectect_name:

    folder_name:

    '''
    # 抓出為【正常模型】的所有檔案名稱
    save = []
    for i in os.listdir():
        if dectect_name in i:
            save.append(i)

    # save=[i for i in os.listdir() if plot_name2 in i]

    # make folder
    ff = [i for i in save if not '.' in i]
    ff = [i for i in ff if '（' in i]

    try:
        os.makedirs(folder_name)
        folder_namenew = folder_name

    except:
        try:
            os.makedirs(folder_name + '（' + str(0) + '）')
            folder_namenew = folder_name + '（' + str(0) + '）'
        except:
            for i in range(0, 10):
                iinn = [j for j in ff if folder_name + '（' + str(i) + '）' in j]
                if len(iinn) == 0:
                    os.makedirs(folder_name + '（' + str(i) + '）')
                    folder_namenew = folder_name + '（' + str(i) + '）'
                    break

                # break

    # move files to that created folder
    import shutil

    save = [i for i in save if '.' in i]
    for m in save:
        shutil.move(m, folder_namenew)


def upload_blob(bucket_name, source_file_name, destination_blob_name, json_os):
    """
    上傳檔案到bucket
    bucket_name = "我們的bucket名稱"
    source_file_name = "我們要上傳檔案的所在「路徑」"
    destination_blob_name = "上傳storage後的名稱要叫做什麼"
    """

    # 環境設定
    os.environ['GOOGLE_APPLICATION_CREDENTIALS'] = json_os

    storage_client = storage.Client()
    bucket = storage_client.bucket(bucket_name)
    blob = bucket.blob(destination_blob_name)

    blob.upload_from_filename(source_file_name)

    print("檔案 {} 已經被上傳到 {}.".format(source_file_name, bucket_name))


def long_speech_recognition_extra_function(
    # 環境參數設定
    json='new.json',
    speech_path=os.getcwd() + '/音檔/06_訪談_小孩教育.wav',
    # 音檔設定
    sample_rate_hertz=48000,
    language_code="cmn-Hant-TW",
    bucket_name='speech2text_howard2',
    destination_blob_name='06_訪談_小孩教育.wav',
    audio_type='WAV',
    # 辨識功能
    enable_word_confidence=True,
    enable_word_time_offsets=True,
    enable_automatic_punctuation=True,
    # 區分說話人員
    enable_speaker_diarization=False,
    audio_channel_count=1,
    diarization_speaker_count=2,
    # 音檔修改專用
    speech_contexts=[''],
    # 產出文件
    enable_docx=True,
    enable_xlsx=True,
    enable_wordcloud=True,
):
    """

    [summary]

    Args:
        json ([json]): [憑證檔]
        speech_path ([string]): [路徑]
        sample_rate_hertz ([int]): [音頻赫茲]
        language_code ([string]): [選擇辨識語音]
        bucket_name ([string]): [雲端storage位置]
        destination_blob_name ([string]): [雲端storage裡面要辨識的檔案]
        enable_word_confidence ([bool]): [True爲開啟逐字信心水準設定]
        enable_word_time_offsets ([bool]): [True爲開啟逐字時間區間設定]
        enable_automatic_punctuation ([bool]): [True爲開啟「標點符號」設定]
        audio_type ([string]): [音檔是什麼類型]
        audio_channel_count ([int]): [聲道數目]
        enable_speaker_diarization ([bool]): [ True爲設定區分不同說話者]
        diarization_speaker_count ([int]): [設定區分多少說話者]
        speech_contexts ([list]): [要改善的字詞]
    """

    # -------------- 前置設定作業 -----------
    # 憑證檔
    json_os = json

    # 環境設定
    os.environ['GOOGLE_APPLICATION_CREDENTIALS'] = json_os

    # 切換GCP客戶端模式
    client = speech_v1p1beta1.SpeechClient()

    # -------------- 長語音轉字檔 -----------
    # 上傳檔案
    upload_blob(
        bucket_name=bucket_name,
        source_file_name=speech_path,
        destination_blob_name=destination_blob_name,
        json_os=json_os,
    )

    import time

    start_time = time.time()

    # 讀取音檔，變成JSON形式
    gcs_uri = "gs://" + bucket_name + '/' + destination_blob_name
    audio = speech_v1p1beta1.RecognitionAudio(uri=gcs_uri)

    # 判斷副檔名，決定編碼格式
    if audio_type == 'MP3':
        encoding = speech.RecognitionConfig.AudioEncoding.ENCODING_UNSPECIFIED
        sample_rate_hertz = sample_rate_hertz

    elif audio_type == 'WAV':
        encoding = speech.RecognitionConfig.AudioEncoding.LINEAR16
        sample_rate_hertz = sample_rate_hertz

    if enable_speaker_diarization:
        language_code = "cmn-Hant-TW"

    # 產出型態設定

    speech_contexts = {"phrases": speech_contexts, "boost": 20}

    config = speech_v1p1beta1.RecognitionConfig(
        audio_channel_count=audio_channel_count,  # 雙聲道
        encoding=encoding,  # 編碼方式
        sample_rate_hertz=sample_rate_hertz,  # 音頻
        language_code=language_code,  # 語言
        # 新功能
        enable_word_confidence=enable_word_confidence,  # 開啟逐字信心水準設定
        enable_automatic_punctuation=enable_automatic_punctuation,  # 開啟標點符號設定
        enable_word_time_offsets=enable_word_time_offsets,  # 開啟逐字時間區間設定
        # 區分人
        enable_speaker_diarization=enable_speaker_diarization,
        diarization_speaker_count=diarization_speaker_count,  # 2個人
        speech_contexts=[speech_contexts],
    )

    # 向GCS請求抓取資料辨識
    operation = client.long_running_recognize(config=config, audio=audio)

    # 轉換辨識結果
    response = operation.result().results

    # 結果呈現
    for result in response:
        print(result.alternatives[0].transcript)  # 辨識結果
        print(result.alternatives[0].confidence)  # 信心水準

    if len(response) == 0:
        return print('注意！沒有聲音')

    print("\n語音辨識秒數 %s  " % (round(time.time() - start_time)))

    # 印出
    print(speech_path)

    # ---------------------- 製作成【逐字信心水準表】DataFrame-----------
    if enable_word_confidence:
        text_list = []
        conf_list = []
        time_list = []

        # 抓取不同句子的words
        for paragraph in response:
            # 將內容轉換為逐字形式
            para = paragraph.alternatives[0].words

            # 在一層迴圈，轉換結果至可讀形式
            for wo in para:
                # 文字
                text_list.append(wo.word)
                # 信心水準
                conf_list.append(wo.confidence)
                # 時間區段
                time_list.append(str(wo.start_time) + '~' + str(wo.end_time))

        word_frame = pd.DataFrame(
            {'斷字': text_list, '機器認字信心水準': conf_list, '始末時間': time_list}
        )

        word_frame['優先查看排序'] = word_frame['機器認字信心水準'].rank(method='min')

        word_frame.to_csv(
            destination_blob_name + '_逐字信心水準表.csv', encoding='utf-8-sig'
        )

    # ---------------------- 製作成製作成【不同談話者區分表】-----------
    if enable_speaker_diarization:
        rec = response[-1].alternatives[0].words

        speaker = []
        text_list = []
        s_time = []
        e_time = []

        for i in rec:
            # 說話者
            speaker.append(i.speaker_tag)
            # 辨識文字
            text_list.append(i.word)
            # print(i.word)
            # 開始時間
            s_time.append(
                i.start_time.seconds + i.start_time.microseconds * 1e-6
            )
            # 結束時間
            e_time.append(i.end_time.seconds + i.end_time.microseconds * 1e-6)

        # 轉為DataFrame
        speaker_frame = pd.DataFrame(
            list(zip(speaker, text_list, s_time, e_time)),
            columns=['說話者編號', '斷字', '開始時間', '結束時間'],
        )

        speaker_frame = translate_chinese(df=speaker_frame, lan="zh_TW")

        speaker_frame['說話者編號'] = '說話者編號' + speaker_frame['說話者編號'].astype(str)

        speaker_frame.to_csv(
            destination_blob_name + '_不同談話者區分表.csv', encoding='utf-8-sig'
        )

        speaker_frame_printdf = speaker_frame.copy()

        # 最終結果：如：皓軒說：「hello大家好」
        result = []

        # 歸納同一個人的話語
        wording = []

        for i in range(0, len(speaker_frame_printdf['說話者編號'])):
            # 如果等於最後一個長度
            if i == (len(speaker_frame_printdf['說話者編號']) - 1):
                print('結束')
                result.append(
                    speaker_frame_printdf['說話者編號'].iloc[i]
                    + '說：「'
                    + ''.join(wording)
                    + '」'
                )
                wording = []

            # 本次的i == 下一個i，也就是還是同一個人的話，就還是歸納成一個人
            elif (
                speaker_frame_printdf['說話者編號'].iloc[i]
                == speaker_frame_printdf['說話者編號'].iloc[i + 1]
            ):
                wording.append(speaker_frame_printdf.iloc[i]['斷字'])

            else:
                # 如果不同人，需要進行歸納
                result.append(
                    speaker_frame_printdf['說話者編號'].iloc[i]
                    + '說：「'
                    + ''.join(wording)
                    + '」'
                )
                wording = []

        wording = pd.DataFrame({'談話': result})
        wording.to_csv(
            destination_blob_name + '_不同談話者講話內容.csv', encoding='utf-8-sig'
        )

    # ---------------------- 三種產出 ------------------------
    # 有分成單音道（單人講話）與多音道（多人講話）

    if audio_channel_count == 1:
        # ---------------------- 三種產出 - 單音道 ------------------------

        # -------文章認字信心矩陣--------

        recog_list = []
        recog_confidence = []
        recog_start = []
        recog_end = []

        for result in response:
            # 抓出文字辨識結果
            text = result.alternatives[0].transcript
            recog_list.append(text)

            # 抓出信賴區間
            conf = result.alternatives[0].confidence
            recog_confidence.append(conf)

            # 抓出開始時間
            s_time = str(result.alternatives[0].words[0].start_time)
            recog_start.append(s_time)

            # 抓出結束時間
            e_time = str(result.alternatives[0].words[-1].end_time)
            recog_end.append(e_time)

        confid_frame = pd.DataFrame(
            list(zip(recog_list, recog_confidence, recog_start, recog_end)),
            columns=['文章段句', '機器認字信心水準', '開始時間', '結束時間'],
        )

        confid_frame['改善順序'] = (
            confid_frame['機器認字信心水準']
            .rank(ascending=True, method='min')
            .astype(int)
        )

        if enable_xlsx:
            confid_frame.to_csv(
                destination_blob_name + '_文章認字信心矩陣.csv', encoding='utf-8-sig'
            )

        # -------製作文字雲--------
        if isinstance(confid_frame, str):
            data = confid_frame.copy()
            print('data is str')

        elif isinstance(confid_frame, pd.DataFrame):
            data_text_list = confid_frame['文章段句'].values.tolist()
            data = ' '.join(data_text_list)
            print('data pd data frame')

        jieba_cut = " ".join(jieba.cut(data))
        with open('stopwords.txt', encoding='UTF-8') as f:
            stopwords = f.readlines()

        stopwords = [w.replace('\n', '') for w in stopwords]  # 換行改成空格
        stopwords = [w.replace(' ', '') for w in stopwords]  # 全型空格改成半型空格
        stopwords.append('\n')  # 新增規則:換行符號
        stopwords.append('\n   \n')  # 新增規則: 兩個換行符號
        stopwords.append('\x0b')  # 新增規則: \x0b

        word_tokens = word_tokenize(jieba_cut)
        filtered_sentence = [w for w in word_tokens if not w in stopwords]

        jieba_cut = ' '.join(filtered_sentence)
        jieba_cut

        if enable_wordcloud:
            wordcloud = WordCloud(
                collocations=False,
                font_path='NotoSansCJKjp-Black.otf',  # 字體設定(是中文一定要設定，否則會是亂碼)
                width=800,  # 圖片寬度
                height=600,  # 圖片高度
                mask=None,
                background_color='white',  # 圖片底色
                margin=2,  # 文字之間的間距
            ).generate(
                jieba_cut
            )  # 要放入的文字
            image = wordcloud.to_image()
            image.show()
            image.save(destination_blob_name + '_文字雲.png')

        # ------- 文章逐字稿 --------
        timer_translist = []
        # 透過迴圈逐一將時間軸取出
        for i in range(0, confid_frame.shape[0]):
            timer_translist.append(
                str(
                    '【'
                    + str(confid_frame.iloc[i, 2])
                    + ' to '
                    + str(confid_frame.iloc[i, 3])
                    + '】'
                )
            )

        document = Document()
        document.styles['Normal'].font.name = u'宋體'  # 設定字體
        document.styles['Normal']._element.rPr.rFonts.set(
            qn('w:eastAsia'), u'宋體'
        )  # 設定字體
        document.add_heading('文章逐字稿', 0)  # 設定標題
        document.add_paragraph(
            '機器認字信心水準'
            + str(round(confid_frame['機器認字信心水準'].mean(), 2))
            + '\n\n'
        )  # 添加段落 － 信心水準

        for i in range(0, confid_frame.shape[0], 1):
            document.add_paragraph(
                confid_frame['文章段句'].iloc[i] + timer_translist[i]
            )

        if enable_wordcloud:
            document.add_picture(
                destination_blob_name + '_文字雲.png', width=Cm(15), height=Cm(13)
            )

        document.save(destination_blob_name + '_文章逐字稿.docx')

    else:
        # ---------------------- 三種產出 - 多音道 ------------------------

        two_record_data = pd.read_csv(destination_blob_name + '_不同談話者講話內容.csv')

        document = Document()
        document.styles['Normal'].font.name = u'宋體'  # 設定字體
        document.styles['Normal']._element.rPr.rFonts.set(
            qn('w:eastAsia'), u'宋體'
        )  # 設定字體
        document.add_heading('文章逐字稿', 0)  # 設定標題

        for i in range(0, two_record_data.shape[0], 1):
            document.add_paragraph(two_record_data['談話'].iloc[i])

        document.save(destination_blob_name + '_多人_文章逐字稿.docx')

    # 音檔秒數計時
    if audio_type == 'WAV':
        from scipy.io import wavfile

        Fs, data = wavfile.read(speech_path)
        t = data.size / Fs
        print('音檔秒數', round(t))

    elif audio_type == 'MP3':
        from mutagen.mp3 import MP3

        audio = MP3(speech_path)
        print('音檔秒數', audio.info.length)

    print('-------------------------------------------')

    return response


def auto_generate_subtitle_NashVersion(word_table, auto_select_range, video):
    auto_select_list = []
    for i in range(0, len(auto_select_range)):
        # 先抓第一個斷句前與後的文字 [0 ~ 11 ~ 17]
        if i == 0:
            # 開始時間
            autotemp = auto_select_range[i]
            temp = word_table[
                (word_table.index >= 0) & (word_table.index < autotemp)
            ]

            # 將該範圍的字詞組合成句子
            word_cut = ''.join(temp.word.tolist())

            # 演練題 - 我如何將 word_cut 這句的起始時間與結束時間抓出來?
            sentence_start_time = temp['start_time'].loc[0]
            sentence_end_time = temp['end_time'].loc[len(temp) - 1]

            '''
                把時間轉換成錄音檔的格式
                因為如果秒數剛好是整數的話後面就不會有微秒
                但錄音檔就算沒有微秒也要顯示000
                因此增加了下述判斷式
            '''

            # 影片特製秒數轉換
            sentence_start_time, sentence_end_time = srt_sec_format(
                sentence_start_time, sentence_end_time
            )

            auto_select_list.append(
                [word_cut, sentence_start_time, sentence_end_time]
            )

        # 第二個後的斷句，僅抓後面的斷字區間 [17 ~25], [25 ~ 32]...
        else:  # i > 0
            # 如果不是最後一個數字，就執行
            # 因為 [208]後面沒有數字了
            if i != len(auto_select_range) - 1:
                # 一次抓兩個出來，例如: [17 ~ 25]
                autotemp = auto_select_range[i - 1 : i + 1]
                temp = word_table[
                    (word_table.index >= autotemp[0])
                    & (word_table.index < autotemp[1])
                ]

                word_cut = ''.join(temp.word.tolist())
                sentence_start_time = temp['start_time'].loc[autotemp[0]]
                sentence_end_time = temp['end_time'].loc[autotemp[1] - 1]

                # 影片特製秒數轉換
                sentence_start_time, sentence_end_time = srt_sec_format(
                    sentence_start_time, sentence_end_time
                )

                auto_select_list.append(
                    [word_cut, sentence_start_time, sentence_end_time]
                )

            # 如果是最後一個數字，就執行else
            else:
                print('last one')
                autotemp = auto_select_range[i]
                temp = word_table[word_table.index >= autotemp]

                word_cut = ''.join(temp.word.tolist())
                sentence_start_time = temp['start_time'].loc[autotemp]
                sentence_end_time = temp['end_time'].loc[
                    autotemp + len(temp) - 1
                ]

                # 影片特製秒數轉換
                sentence_start_time, sentence_end_time = srt_sec_format(
                    sentence_start_time, sentence_end_time
                )

                auto_select_list.append(
                    [word_cut, sentence_start_time, sentence_end_time]
                )
    # srt_list完成字幕list
    srt_list = []
    i = 1

    # 透過迴圈，逐行新增srt形式字幕
    for write in auto_select_list:
        lines = (
            str(i)
            + '\n'
            + write[1]
            + ' --> '
            + write[2]
            + '\n'
            + write[0]
            + '\n\n'
        )
        lines = lines.replace('拍的', 'Python')
        lines = lines.replace('拍等', 'Python')
        lines = lines.replace('拍本', 'Python')
        lines = lines.replace('拍得', 'Python')
        lines = lines.replace('拍燈', 'Python')
        lines = lines.replace('python', 'Python')
        lines = lines.replace('張傑', '章節')
        lines = lines.replace('panda', 'Panda')
        lines = lines.replace('dataframe', 'DataFrame')
        lines = lines.replace('生理', 'streamlit')
        srt_list.append(lines)
        i += 1

    with open(video + '.srt', 'w', encoding='UTF-8') as f:
        f.write(''.join(srt_list))

    return srt_list
